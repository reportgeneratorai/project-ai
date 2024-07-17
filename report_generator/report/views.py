from django.shortcuts import render
from django.shortcuts import render, HttpResponse, redirect, get_object_or_404
from report.models import Prompt
from report.models import Docx_file
from django.utils import timezone
from django.core.files import File


import os
import docx
import google.generativeai as genai
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from dotenv import load_dotenv
from requests import get
from bs4 import BeautifulSoup
# Create your views here.


def home(request):
    if request.method == 'POST':
        prompt_input = request.POST.get('prompt')
        load_dotenv()

        # Initialize Gemini API
        genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
        generation_config = {
            "temperature": 1,
            "top_p": 0.95,
            "top_k": 64,
            "max_output_tokens": 4000,
            "response_mime_type": "text/plain",
        }
        model = genai.GenerativeModel(
            model_name="gemini-1.5-flash",
            generation_config=generation_config,
            system_instruction="You are a chat bot which is used to generate Projects report of huge paragraphs on given topic, your response should be proper and reliable for storing in a word file in proper format of project report. Use Heading 1 for main sections and Heading 2 for subheadings."
        )

        def generate_report(title):
            document = Document()

            # Title
            title_heading = document.add_heading(title, level=1)
            title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Content from Gemini API
            content = fetch_content(title)
            if content:
                paragraphs = process_content(document, content)
                images = fetch_images(title)

                # Insert images at suitable places within the paragraphs
                insert_images(document, paragraphs, images)
            else:
                document.add_paragraph("No content available for this topic.")

            # Save the document
            temp_dir = 'tmp'
            os.makedirs(temp_dir, exist_ok=True)
            
            filepath = os.path.join(temp_dir, f"{title}.docx")
            document.save(filepath)
            return filepath

        def fetch_content(title):
            try:
                response = model.generate_content(f"Generate a detailed and professional Micro Project Report on {title} with proper structure, suitable for engineering students. Include sections such as Introduction, Working Principle, Methodology, Classification, Applications, Results, Conclusion, and References.")
                
                if response and response.text:
                    return response.text
                else:
                    return "No content available for this topic."
            except Exception as e:
                print(f"Error fetching content: {e}")
            return "No content available for this topic."

        def process_content(document, content):
            paragraphs = []
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith("## "):
                    heading = document.add_heading(line[3:], level=1)
                    for run in heading.runs:
                        run.font.size = Pt(18)
                        run.font.name = 'Arial'
                elif line.startswith("### "):
                    heading = document.add_heading(line[4:], level=2)
                    for run in heading.runs:
                        run.font.size = Pt(16)
                        run.font.name = 'Arial'
                elif line.startswith("* "):
                    heading = document.add_heading(line[2:], level=3)
                    for run in heading.runs:
                        run.font.size = Pt(14)
                        run.font.name = 'Arial'
                else:
                    p = document.add_paragraph()
                    parts = line.split("**")
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:
                            run.bold = True
                        run.font.size = Pt(12)
                        run.font.name = 'Arial'
                    paragraphs.append(p)
            return paragraphs

        def fetch_images(title):
            try:
                response = get(f"https://www.google.com/search?tbm=isch&q={title}")
                if response.status_code == 200:
                    soup = BeautifulSoup(response.text, 'html.parser')
                    img_tags = soup.find_all('img')
                    img_urls = []
                    for img in img_tags:
                        img_url = img.get('src')
                        if img_url and img_url.startswith('http'):
                            img_urls.append(img_url)
                    return img_urls[:5]  # Return the first 5 valid image URLs
            except Exception as e:
                print(f"Error fetching images: {e}")
            return []

        def insert_images(document, paragraphs, images):
            tmp_dir = 'tmp'
            os.makedirs(tmp_dir, exist_ok=True)  # Ensure the tmp directory exists

            for i, paragraph in enumerate(paragraphs):
                if i % 5 == 0 and i // 5 < len(images):
                    img_url = images[i // 5]
                    response = get(img_url)
                    if response.status_code == 200:
                        image_path = os.path.join(tmp_dir, f'image_{i // 5}.jpg')
                        with open(image_path, 'wb') as img_file:
                            img_file.write(response.content)
                        run = paragraph.add_run()
                        run.add_break()
                        run.add_picture(image_path, width=Inches(4.0))
                        os.remove(image_path)  # Clean up image after adding

        # Generate the report
        filepath = generate_report(prompt_input)

        # Save the file in the Docx_file model with session key
        with open(filepath, 'rb') as f:
            thefile = Docx_file(file=File(f), session_key=request.session.session_key)
            thefile.save()

    # Retrieve the file for the current session
    docx_files = Docx_file.objects.filter(session_key=request.session.session_key)
    if docx_files.exists():
        file_url = docx_files.last().file.url
    else:
        file_url = None

    return render(request, 'home/index.html', {'file_url': file_url})

'''if request.method == 'POST':
            prompt_input = request.POST.get('prompt')
            if prompt_input:
                prompt = Prompt(p_input=prompt_input, timeStamp=timezone.now())
                prompt.save()
                return redirect('home')  # Redirect to the home page after saving '''
        
        
